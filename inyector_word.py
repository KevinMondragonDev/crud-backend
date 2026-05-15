import json
from docx import Document


def reemplazar_texto_en_word(ruta_plantilla, ruta_salida, archivo_json):
    """Lee un JSON generado por la IA y reemplaza los marcadores {{CLAVE}} en la plantilla .docx."""
    with open(archivo_json, 'r', encoding='utf-8') as f:
        datos = json.load(f)

    doc = Document(ruta_plantilla)

    for p in doc.paragraphs:
        for key, value in datos.items():
            marcador = f"{{{{{key}}}}}"
            if marcador in p.text:
                p.text = p.text.replace(marcador, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in datos.items():
                    marcador = f"{{{{{key}}}}}"
                    if marcador in cell.text:
                        cell.text = cell.text.replace(marcador, str(value))

    doc.save(ruta_salida)
    print(f"Documento generado con éxito: {ruta_salida}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) == 4:
        reemplazar_texto_en_word(sys.argv[1], sys.argv[2], sys.argv[3])
    else:
        print("Uso: python inyector_word.py <plantilla.docx> <salida.docx> <datos.json>")
