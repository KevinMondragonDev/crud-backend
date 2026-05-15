"""Script para generar las plantillas DDR y DDS con marcadores {{CLAVE}}."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def estilo_encabezado(paragraph, texto, nivel=1):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(texto)
    run.bold = True
    run.font.size = Pt(16 if nivel == 1 else 13)
    run.font.color.rgb = RGBColor(0, 51, 102)


def agregar_tabla(doc, filas, columnas, datos_celdas):
    """Crea una tabla con datos_celdas como lista de listas de strings."""
    table = doc.add_table(rows=filas, cols=columnas)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, fila in enumerate(datos_celdas):
        for j, texto in enumerate(fila):
            cell = table.rows[i].cells[j]
            cell.text = texto
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    return table


def generar_plantilla_ddr(ruta_salida):
    doc = Document()

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEFINICIÓN DE REQUERIMIENTOS (DDR)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    info_portada = [
        ["Campo", "Valor"],
        ["Proyecto", "{{PROYECTO_NOMBRE}}"],
        ["Versión", "{{VERSION}}"],
        ["Fecha", "{{FECHA}}"],
        ["Autor", "{{AUTOR}}"],
        ["Estado", "{{ESTADO}}"],
    ]
    agregar_tabla(doc, len(info_portada), 2, info_portada)

    doc.add_page_break()

    # 1. Introducción
    p = doc.add_paragraph()
    estilo_encabezado(p, "1. Introducción")
    doc.add_paragraph("{{INTRODUCCION}}")

    # 1.1 Propósito
    p = doc.add_paragraph()
    estilo_encabezado(p, "1.1 Propósito", 2)
    doc.add_paragraph("{{PROPOSITO}}")

    # 1.2 Alcance
    p = doc.add_paragraph()
    estilo_encabezado(p, "1.2 Alcance", 2)
    doc.add_paragraph("{{ALCANCE}}")

    # 2. Descripción General del Sistema
    p = doc.add_paragraph()
    estilo_encabezado(p, "2. Descripción General del Sistema")
    doc.add_paragraph("{{DESCRIPCION_GENERAL}}")

    # 2.1 Stack Tecnológico
    p = doc.add_paragraph()
    estilo_encabezado(p, "2.1 Stack Tecnológico", 2)
    doc.add_paragraph("{{STACK_TECNOLOGICO}}")

    # 3. Requerimientos Funcionales
    p = doc.add_paragraph()
    estilo_encabezado(p, "3. Requerimientos Funcionales")

    req_func = [
        ["ID", "Descripción", "Prioridad", "Fuente en Código"],
        ["{{REQ001_ID}}", "{{REQ001_DESC}}", "{{REQ001_PRIORIDAD}}", "{{REQ001_FUENTE}}"],
        ["{{REQ002_ID}}", "{{REQ002_DESC}}", "{{REQ002_PRIORIDAD}}", "{{REQ002_FUENTE}}"],
        ["{{REQ003_ID}}", "{{REQ003_DESC}}", "{{REQ003_PRIORIDAD}}", "{{REQ003_FUENTE}}"],
        ["{{REQ004_ID}}", "{{REQ004_DESC}}", "{{REQ004_PRIORIDAD}}", "{{REQ004_FUENTE}}"],
        ["{{REQ005_ID}}", "{{REQ005_DESC}}", "{{REQ005_PRIORIDAD}}", "{{REQ005_FUENTE}}"],
    ]
    agregar_tabla(doc, len(req_func), 4, req_func)

    # 4. Requerimientos No Funcionales
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "4. Requerimientos No Funcionales")

    req_no_func = [
        ["ID", "Descripción", "Categoría"],
        ["{{RNF001_ID}}", "{{RNF001_DESC}}", "{{RNF001_CATEGORIA}}"],
        ["{{RNF002_ID}}", "{{RNF002_DESC}}", "{{RNF002_CATEGORIA}}"],
        ["{{RNF003_ID}}", "{{RNF003_DESC}}", "{{RNF003_CATEGORIA}}"],
    ]
    agregar_tabla(doc, len(req_no_func), 3, req_no_func)

    # 5. Matriz de Pruebas
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "5. Matriz de Pruebas")

    # 5.1 Escenarios de Prueba
    p = doc.add_paragraph()
    estilo_encabezado(p, "5.1 Escenarios de Prueba", 2)

    escenarios = [
        ["ID Escenario", "Descripción del Escenario", "Precondiciones", "Resultado Esperado"],
        ["{{ESC001_ID}}", "{{ESC001_DESC}}", "{{ESC001_PRECONDICIONES}}", "{{ESC001_RESULTADO}}"],
        ["{{ESC002_ID}}", "{{ESC002_DESC}}", "{{ESC002_PRECONDICIONES}}", "{{ESC002_RESULTADO}}"],
        ["{{ESC003_ID}}", "{{ESC003_DESC}}", "{{ESC003_PRECONDICIONES}}", "{{ESC003_RESULTADO}}"],
    ]
    agregar_tabla(doc, len(escenarios), 4, escenarios)

    # 5.2 Casos de Prueba
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "5.2 Casos de Prueba", 2)

    casos = [
        ["ID Caso", "Escenario Asociado", "Datos de Entrada", "Pasos", "Resultado Esperado", "Resultado Obtenido"],
        ["{{CP001_ID}}", "{{CP001_ESCENARIO}}", "{{CP001_DATOS}}", "{{CP001_PASOS}}", "{{CP001_ESPERADO}}", "{{CP001_OBTENIDO}}"],
        ["{{CP002_ID}}", "{{CP002_ESCENARIO}}", "{{CP002_DATOS}}", "{{CP002_PASOS}}", "{{CP002_ESPERADO}}", "{{CP002_OBTENIDO}}"],
        ["{{CP003_ID}}", "{{CP003_ESCENARIO}}", "{{CP003_DATOS}}", "{{CP003_PASOS}}", "{{CP003_ESPERADO}}", "{{CP003_OBTENIDO}}"],
    ]
    agregar_tabla(doc, len(casos), 6, casos)

    doc.save(ruta_salida)
    print(f"Plantilla DDR generada: {ruta_salida}")


def generar_plantilla_dds(ruta_salida):
    doc = Document()

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DISEÑO DEL SISTEMA (DDS)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    info_portada = [
        ["Campo", "Valor"],
        ["Proyecto", "{{PROYECTO_NOMBRE}}"],
        ["Versión", "{{VERSION}}"],
        ["Fecha", "{{FECHA}}"],
        ["Autor", "{{AUTOR}}"],
        ["Estado", "{{ESTADO}}"],
    ]
    agregar_tabla(doc, len(info_portada), 2, info_portada)

    doc.add_page_break()

    # 1. Introducción
    p = doc.add_paragraph()
    estilo_encabezado(p, "1. Introducción")
    doc.add_paragraph("{{DDS_INTRODUCCION}}")

    # 2. Arquitectura del Sistema
    p = doc.add_paragraph()
    estilo_encabezado(p, "2. Arquitectura del Sistema")
    doc.add_paragraph("{{ARQUITECTURA}}")

    # 2.1 Diagrama de Componentes
    p = doc.add_paragraph()
    estilo_encabezado(p, "2.1 Diagrama de Componentes (Textual)", 2)
    doc.add_paragraph("{{DIAGRAMA_COMPONENTES}}")

    # 3. Diseño de Módulos
    p = doc.add_paragraph()
    estilo_encabezado(p, "3. Diseño de Módulos")

    modulos = [
        ["Módulo", "Clase/Archivo", "Responsabilidad", "Dependencias"],
        ["{{MOD001_NOMBRE}}", "{{MOD001_CLASE}}", "{{MOD001_RESPONSABILIDAD}}", "{{MOD001_DEPENDENCIAS}}"],
        ["{{MOD002_NOMBRE}}", "{{MOD002_CLASE}}", "{{MOD002_RESPONSABILIDAD}}", "{{MOD002_DEPENDENCIAS}}"],
        ["{{MOD003_NOMBRE}}", "{{MOD003_CLASE}}", "{{MOD003_RESPONSABILIDAD}}", "{{MOD003_DEPENDENCIAS}}"],
    ]
    agregar_tabla(doc, len(modulos), 4, modulos)

    # 4. Diseño Detallado de Clases
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "4. Diseño Detallado de Clases")

    clases = [
        ["Clase", "Paquete", "Métodos Principales", "Atributos", "Tipo"],
        ["{{CLS001_NOMBRE}}", "{{CLS001_PAQUETE}}", "{{CLS001_METODOS}}", "{{CLS001_ATRIBUTOS}}", "{{CLS001_TIPO}}"],
        ["{{CLS002_NOMBRE}}", "{{CLS002_PAQUETE}}", "{{CLS002_METODOS}}", "{{CLS002_ATRIBUTOS}}", "{{CLS002_TIPO}}"],
        ["{{CLS003_NOMBRE}}", "{{CLS003_PAQUETE}}", "{{CLS003_METODOS}}", "{{CLS003_ATRIBUTOS}}", "{{CLS003_TIPO}}"],
    ]
    agregar_tabla(doc, len(clases), 5, clases)

    # 5. Flujo de Datos
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "5. Flujo de Datos")
    doc.add_paragraph("{{FLUJO_DATOS}}")

    # 6. Manejo de Entrada/Salida
    p = doc.add_paragraph()
    estilo_encabezado(p, "6. Manejo de Entrada/Salida")

    io_table = [
        ["Tipo", "Descripción", "Validación", "Clase Responsable"],
        ["{{IO001_TIPO}}", "{{IO001_DESC}}", "{{IO001_VALIDACION}}", "{{IO001_CLASE}}"],
        ["{{IO002_TIPO}}", "{{IO002_DESC}}", "{{IO002_VALIDACION}}", "{{IO002_CLASE}}"],
        ["{{IO003_TIPO}}", "{{IO003_DESC}}", "{{IO003_VALIDACION}}", "{{IO003_CLASE}}"],
    ]
    agregar_tabla(doc, len(io_table), 4, io_table)

    # 7. Decisiones de Diseño
    doc.add_paragraph()
    p = doc.add_paragraph()
    estilo_encabezado(p, "7. Decisiones de Diseño")
    doc.add_paragraph("{{DECISIONES_DISENO}}")

    doc.save(ruta_salida)
    print(f"Plantilla DDS generada: {ruta_salida}")


if __name__ == "__main__":
    generar_plantilla_ddr("DDR_plantilla.docx")
    generar_plantilla_dds("DDS_plantilla.docx")
